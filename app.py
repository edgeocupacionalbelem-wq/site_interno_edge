from flask import Flask, render_template, request, send_file, flash, redirect, url_for, jsonify
import os, re, tempfile, unicodedata, zipfile, io, sqlite3, shutil, subprocess
from io import BytesIO
from pathlib import Path
from datetime import datetime
from difflib import SequenceMatcher

import pandas as pd
from werkzeug.utils import secure_filename
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from docxtpl import DocxTemplate, RichText
from docx import Document

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "sistema-interno-4-funcionalidades")

APP_TITLE = "Sistema Interno"
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "ENCAMINHAMENTOS PERIODICO MCP.docx")
ALLOWED_EXTENSIONS = {".xls", ".xlsx", ".html", ".htm"}
RENUM_ALLOWED_EXTENSIONS = {".docx", ".zip"}

FISICO_TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "ATESTADO_FISICO_MENTAL_TEMPLATE.docx")
BASE_DIR = os.path.dirname(__file__)
DATA_DIR = os.environ.get("RENDER_DISK_PATH") or os.environ.get("DATA_DIR") or BASE_DIR
os.makedirs(DATA_DIR, exist_ok=True)
FISICO_DB_PATH = os.path.join(DATA_DIR, "fisico_mental.db")

# =========================
# UTILIDADES GERAIS
# =========================
def normalize_text(value) -> str:
    if pd.isna(value):
        return ""
    text = str(value).strip().upper()
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    return re.sub(r"\s+", " ", text)

def sanitize_filename(name: str) -> str:
    name = re.sub(r'[\\/:*?"<>|]+', "_", str(name))
    name = re.sub(r"\s+", " ", name).strip()
    return name or "arquivo"

def unique_path(path: str) -> str:
    base, ext = os.path.splitext(path)
    if not os.path.exists(path):
        return path
    i = 2
    while True:
        new = f"{base} ({i}){ext}"
        if not os.path.exists(new):
            return new
        i += 1

def encontrar_coluna(df, candidatos, obrigatoria=False):
    cols = list(df.columns)
    normalizadas = {str(c).strip().lower(): c for c in cols}
    for nome in candidatos:
        key = nome.strip().lower()
        if key in normalizadas:
            return normalizadas[key]
    for c in cols:
        c_norm = str(c).strip().lower()
        for nome in candidatos:
            if nome.strip().lower() in c_norm:
                return c
    if obrigatoria:
        raise ValueError(f"Coluna não encontrada. Esperado um destes nomes: {', '.join(candidatos)}")
    return None

def limpar_nome_arquivo(nome_arquivo):
    nome = os.path.splitext(os.path.basename(str(nome_arquivo)))[0]
    nome = re.sub(r"\(\d+\)$", "", nome).strip()
    nome = re.sub(r"\s{2,}", " ", nome)
    return nome

def limpar_nome_pasta_arquivo(nome):
    nome = str(nome or "")
    for c in r'\/:*?"<>|':
        nome = nome.replace(c, "")
    nome = re.sub(r"\s{2,}", " ", nome).strip()
    return nome

def somente_numeros(texto):
    return re.sub(r"\D", "", str(texto or ""))

def formatar_documento(doc):
    numeros = somente_numeros(doc)
    if len(numeros) == 14:
        return f"{numeros[:2]}.{numeros[2:5]}.{numeros[5:8]}/{numeros[8:12]}-{numeros[12:]}"
    if len(numeros) == 11:
        return f"{numeros[:3]}.{numeros[3:6]}.{numeros[6:9]}-{numeros[9:]}"
    return ""

def extrair_documento_do_final_do_arquivo(nome_arquivo):
    nome = limpar_nome_arquivo(nome_arquivo)
    partes = [p.strip() for p in nome.split(" - ") if p.strip()]
    candidatos = []
    if partes:
        candidatos.append(partes[-1])
    m = re.search(r'([0-9.\-\/]+)\s*$', nome)
    if m:
        candidatos.append(m.group(1).strip())
    for cand in candidatos:
        numeros = somente_numeros(cand)
        if len(numeros) == 14:
            return formatar_documento(numeros)
        if len(numeros) == 11:
            return formatar_documento(numeros)
        encontrados = re.findall(r'\d+', cand)
        if encontrados:
            juntos = "".join(encontrados)
            if len(juntos) >= 14:
                return formatar_documento(juntos[-14:])
            if len(juntos) >= 11:
                return formatar_documento(juntos[-11:])
    nums = re.findall(r'\d+', nome)
    if nums:
        juntos = "".join(nums)
        if len(juntos) >= 14:
            return formatar_documento(juntos[-14:])
        if len(juntos) >= 11:
            return formatar_documento(juntos[-11:])
    return ""

def nome_mes(m):
    meses = ["", "JANEIRO","FEVEREIRO","MARÇO","ABRIL","MAIO","JUNHO",
             "JULHO","AGOSTO","SETEMBRO","OUTUBRO","NOVEMBRO","DEZEMBRO"]
    return meses[m]

class UploadedMemoryFile(BytesIO):
    def __init__(self, filename, data):
        super().__init__(data)
        self.filename = filename

# =========================
# RELATÓRIOS + BASE DO MÊS
# =========================
def nome_empresa_da_planilha(row, col_empresa, col_setor, nome_arquivo):
    if col_empresa:
        valor = row[col_empresa]
        if not pd.isna(valor) and str(valor).strip():
            return str(valor).strip()
    if col_setor:
        valor = row[col_setor]
        if not pd.isna(valor) and str(valor).strip():
            return str(valor).strip()
    return limpar_nome_arquivo(nome_arquivo)

def criar_relatorio(files, mes):
    wb = Workbook()
    ws = wb.active
    ws.title = "Relatório"

    cor_empresa = PatternFill(start_color="BDD7EE", end_color="BDD7EE", fill_type="solid")
    borda = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    center = Alignment(horizontal="center", vertical="center")
    left_wrap = Alignment(horizontal="left", vertical="center", wrap_text=True)

    linha = 1

    for file in files:
        try:
            file.seek(0)
            df = pd.read_excel(file)

            col_data = encontrar_coluna(df, ["admissao", "admissão", "data"], obrigatoria=True)
            col_nome = encontrar_coluna(df, ["nome", "funcionario", "funcionário"], obrigatoria=True)
            col_empresa = encontrar_coluna(df, ["empresa"])
            col_setor = encontrar_coluna(df, ["setor"])

            df[col_data] = pd.to_datetime(df[col_data], dayfirst=True, errors="coerce")
            df = df.dropna(subset=[col_data, col_nome])
            filtrado = df[df[col_data].dt.month == mes].copy()

            if not filtrado.empty:
                titulo = nome_empresa_da_planilha(filtrado.iloc[0], col_empresa, col_setor, file.filename)
            elif not df.empty:
                titulo = nome_empresa_da_planilha(df.iloc[0], col_empresa, col_setor, file.filename)
            else:
                titulo = limpar_nome_arquivo(file.filename)

            ws.merge_cells(start_row=linha, start_column=1, end_row=linha, end_column=3)
            cell = ws.cell(row=linha, column=1, value=titulo)
            cell.font = Font(size=13, bold=True)
            cell.fill = cor_empresa
            cell.alignment = left_wrap
            for col in range(1, 4):
                ws.cell(row=linha, column=col).border = borda
            ws.row_dimensions[linha].height = 35
            linha += 1

            ws.merge_cells(start_row=linha, start_column=1, end_row=linha, end_column=3)
            cell = ws.cell(row=linha, column=1, value="NOME DO FUNCIONÁRIO")
            cell.font = Font(bold=True)
            cell.alignment = center
            for col in range(1, 4):
                ws.cell(row=linha, column=col).border = borda
            linha += 1

            if filtrado.empty:
                ws.merge_cells(start_row=linha, start_column=1, end_row=linha, end_column=3)
                cell = ws.cell(row=linha, column=1, value=f"NÃO HÁ EXAMES PERIÓDICOS PARA O MÊS DE {nome_mes(mes)}")
                cell.font = Font(color="FF0000", bold=True)
                cell.alignment = center
                for col in range(1, 4):
                    ws.cell(row=linha, column=col).border = borda
                linha += 3
            else:
                for _, row in filtrado.iterrows():
                    nome = "" if pd.isna(row[col_nome]) else str(row[col_nome]).strip()
                    ws.merge_cells(start_row=linha, start_column=1, end_row=linha, end_column=3)
                    cell = ws.cell(row=linha, column=1, value=nome)
                    cell.alignment = Alignment(horizontal="left", vertical="center")
                    for col in range(1, 4):
                        ws.cell(row=linha, column=col).border = borda
                    linha += 1
                linha += 2
        except Exception as e:
            titulo = limpar_nome_arquivo(file.filename)
            ws.merge_cells(start_row=linha, start_column=1, end_row=linha, end_column=3)
            cell = ws.cell(row=linha, column=1, value=titulo)
            cell.font = Font(size=13, bold=True)
            cell.fill = cor_empresa
            cell.alignment = left_wrap
            for col in range(1, 4):
                ws.cell(row=linha, column=col).border = borda
            ws.row_dimensions[linha].height = 35
            linha += 1
            ws.merge_cells(start_row=linha, start_column=1, end_row=linha, end_column=3)
            cell = ws.cell(row=linha, column=1, value="NOME DO FUNCIONÁRIO")
            cell.font = Font(bold=True)
            cell.alignment = center
            for col in range(1, 4):
                ws.cell(row=linha, column=col).border = borda
            linha += 1
            ws.merge_cells(start_row=linha, start_column=1, end_row=linha, end_column=3)
            cell = ws.cell(row=linha, column=1, value=f"Erro ao processar arquivo: {e}")
            cell.font = Font(color="FF0000", bold=True)
            cell.alignment = center
            for col in range(1, 4):
                ws.cell(row=linha, column=col).border = borda
            linha += 3

    ws.column_dimensions["A"].width = 55
    ws.column_dimensions["B"].width = 12
    ws.column_dimensions["C"].width = 12
    return wb

def criar_base(files, mes):
    wb = Workbook()
    ws = wb.active
    ws.title = "Base do Mês"

    borda = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    cabecalho_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    cabecalho_font = Font(bold=True)
    center = Alignment(horizontal="center", vertical="center")

    headers = ["empresa", "cnpj", "nome", "cargo", "Complementares"]
    for idx, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=idx, value=h)
        cell.font = cabecalho_font
        cell.fill = cabecalho_fill
        cell.alignment = center
        cell.border = borda

    linha = 2

    for file in files:
        try:
            file.seek(0)
            df = pd.read_excel(file)

            col_data = encontrar_coluna(df, ["admissao", "admissão", "data"], obrigatoria=True)
            col_nome = encontrar_coluna(df, ["nome", "funcionario", "funcionário"], obrigatoria=True)
            col_cargo = encontrar_coluna(df, ["cargo", "função", "funcao"])
            col_empresa = encontrar_coluna(df, ["empresa"])
            col_setor = encontrar_coluna(df, ["setor"])
            col_comp = encontrar_coluna(df, ["complementares", "complementar"])

            df[col_data] = pd.to_datetime(df[col_data], dayfirst=True, errors="coerce")
            df = df.dropna(subset=[col_data, col_nome])
            filtrado = df[df[col_data].dt.month == mes].copy()

            documento_arquivo = extrair_documento_do_final_do_arquivo(file.filename)

            for _, row in filtrado.iterrows():
                empresa = nome_empresa_da_planilha(row, col_empresa, col_setor, file.filename)
                nome = "" if pd.isna(row[col_nome]) else str(row[col_nome]).strip()
                cargo = "" if not col_cargo or pd.isna(row[col_cargo]) else str(row[col_cargo]).strip()
                complementares = "" if not col_comp or pd.isna(row[col_comp]) else str(row[col_comp]).strip()

                valores = [empresa, documento_arquivo, nome, cargo, complementares]
                for idx, valor in enumerate(valores, start=1):
                    cell = ws.cell(row=linha, column=idx, value=str(valor) if valor is not None else "")
                    cell.border = borda
                    if idx == 2:
                        cell.number_format = "@"
                linha += 1
        except Exception:
            continue

    ws.column_dimensions["A"].width = 45
    ws.column_dimensions["B"].width = 22
    ws.column_dimensions["C"].width = 38
    ws.column_dimensions["D"].width = 28
    ws.column_dimensions["E"].width = 40
    return wb

# =========================
# ENCAMINHAMENTOS
# =========================
def quebrar_complementares(texto):
    if texto is None or (isinstance(texto, float) and pd.isna(texto)):
        return []
    texto = str(texto).strip()
    if not texto or texto.lower() == "nan":
        return []
    return [x.strip() for x in re.split(r";+", texto) if x.strip()]

def gerar_encaminhamentos(file):
    file.seek(0)
    df = pd.read_excel(file)

    col_empresa = encontrar_coluna(df, ["empresa"], obrigatoria=True)
    col_cnpj = encontrar_coluna(df, ["cnpj", "cpf"])
    col_nome = encontrar_coluna(df, ["funcionario", "funcionário", "nome"], obrigatoria=True)
    col_funcao = encontrar_coluna(df, ["funcao", "função", "cargo"])
    col_comp = encontrar_coluna(df, ["complementares"], obrigatoria=False)

    temp_dir = tempfile.mkdtemp()
    out_root = os.path.join(temp_dir, "encaminhamentos")
    os.makedirs(out_root, exist_ok=True)

    for _, row in df.iterrows():
        empresa = "" if pd.isna(row[col_empresa]) else str(row[col_empresa]).strip()
        funcionario = "" if pd.isna(row[col_nome]) else str(row[col_nome]).strip()
        cnpj = "" if not col_cnpj or pd.isna(row[col_cnpj]) else str(row[col_cnpj]).strip()
        funcao = "" if not col_funcao or pd.isna(row[col_funcao]) else str(row[col_funcao]).strip()
        complementares_txt = "" if not col_comp or pd.isna(row[col_comp]) else str(row[col_comp]).strip()

        lista = quebrar_complementares(complementares_txt)
        comps = {f"comp{i+1}": lista[i] if i < len(lista) else "" for i in range(9)}

        pasta_empresa = os.path.join(out_root, limpar_nome_pasta_arquivo(empresa or "SEM EMPRESA"))
        os.makedirs(pasta_empresa, exist_ok=True)

        contexto = {"empresa": empresa, "cnpj": cnpj, "funcionario": funcionario, "funcao": funcao, **comps}

        template = DocxTemplate(TEMPLATE_PATH)
        template.render(contexto)

        nome_base = f"ENCAMINHAMENTO {limpar_nome_pasta_arquivo(funcionario or 'SEM NOME')}.docx"
        destino = os.path.join(pasta_empresa, nome_base)
        contador = 1
        while os.path.exists(destino):
            nome_base = f"ENCAMINHAMENTO {limpar_nome_pasta_arquivo(funcionario or 'SEM NOME')} ({contador}).docx"
            destino = os.path.join(pasta_empresa, nome_base)
            contador += 1
        template.save(destino)

    zip_path = os.path.join(temp_dir, "encaminhamentos.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for root, _, files in os.walk(out_root):
            for name in files:
                full = os.path.join(root, name)
                rel = os.path.relpath(full, out_root)
                z.write(full, rel)
    return zip_path

# =========================
# RENUMERADOR
# =========================
def allowed_renum_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in RENUM_ALLOWED_EXTENSIONS

def encontrar_ultimo_numero(doc: Document) -> str:
    numeros = []
    for i in range(len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip() == "NOTA DE BALCÃO":
            if i + 1 < len(doc.paragraphs):
                num = doc.paragraphs[i + 1].text.strip()
                if num.isdigit():
                    numeros.append(num)
    if numeros:
        return max(numeros, key=lambda x: int(x))
    return "0"

def atualizar_data_documento(doc: Document, nova_data: str) -> int:
    alteradas = 0
    padrao = re.compile(r"^\s*Data:\s*\d{2}/\d{2}/\d{4}\s*$")
    for p in doc.paragraphs:
        texto = p.text.strip()
        if padrao.match(texto):
            novo_texto = f"Data: {nova_data}"
            if p.runs:
                p.runs[0].text = novo_texto
                p.runs[0].bold = True
                for run in p.runs[1:]:
                    run.text = ""
            else:
                p.text = novo_texto
                if p.runs:
                    p.runs[0].bold = True
            alteradas += 1
    return alteradas

def renumerar_documento(caminho_entrada: str, caminho_saida: str, nova_data: str):
    doc = Document(caminho_entrada)
    ultimo = encontrar_ultimo_numero(doc)
    tamanho = len(ultimo)
    numero_atual = int(ultimo) + 1
    i = 0
    alterados = 0
    while i < len(doc.paragraphs):
        if doc.paragraphs[i].text.strip() == "NOTA DE BALCÃO":
            for run in doc.paragraphs[i].runs:
                run.bold = False
            if i + 1 < len(doc.paragraphs):
                texto_num = doc.paragraphs[i + 1].text.strip()
                if texto_num.isdigit():
                    novo_num = str(numero_atual).zfill(tamanho)
                    p_num = doc.paragraphs[i + 1]
                    if p_num.runs:
                        p_num.runs[0].text = novo_num
                        p_num.runs[0].bold = True
                        for run in p_num.runs[1:]:
                            run.text = ""
                    else:
                        p_num.text = novo_num
                        if p_num.runs:
                            p_num.runs[0].bold = True
                    numero_atual += 1
                    alterados += 1
                    i += 1
        i += 1
    datas_alteradas = atualizar_data_documento(doc, nova_data)
    Path(caminho_saida).parent.mkdir(parents=True, exist_ok=True)
    doc.save(caminho_saida)
    return alterados, ultimo, datas_alteradas

# =========================
# E-SOCIAL
# =========================
def normalize_company_name(value) -> str:
    text = normalize_text(value)
    text = re.sub(r"[^\w\s]", " ", text)
    text = re.sub(r"\b(LTDA|EIRELI|ME|EPP|S A|SA|S/S|SS|MATRIZ|FILIAL)\b", " ", text)
    return re.sub(r"\s+", " ", text).strip()

def is_allowed_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS

def extract_cnpj(text: str) -> str:
    digits = re.sub(r"\D", "", ("" if text is None else str(text)))
    return digits[:14] if len(digits) >= 14 else ""

def score_dataframe(df: pd.DataFrame) -> int:
    score = 0
    cols = [normalize_text(c) for c in df.columns]
    for wanted in ["FUNCIONARIO","TIPO DE EXAME","DEPOSITANTE","SETOR","NOME","TIPO","EMPRESA"]:
        if wanted in cols:
            score += 10
    score += min(len(df), 50)
    return score

def list_sheets(path: str):
    suffix = Path(path).suffix.lower()
    if suffix in {".html", ".htm"}:
        return ["Planilha principal"]
    try:
        xl = pd.ExcelFile(path)
        if xl.sheet_names:
            return xl.sheet_names
    except Exception:
        pass
    try:
        tables = pd.read_html(path)
        if tables:
            return ["Planilha principal"]
    except Exception:
        pass
    return ["Planilha principal"]

def read_spreadsheet(path: str, selected_sheet: str | None = None) -> pd.DataFrame:
    suffix = Path(path).suffix.lower()
    if selected_sheet and selected_sheet != "Planilha principal":
        try:
            return pd.read_excel(path, sheet_name=selected_sheet)
        except Exception as exc:
            raise RuntimeError(f"Não foi possível ler a aba '{selected_sheet}' do arquivo {os.path.basename(path)}. Erro: {exc}") from exc
    if suffix == ".xls":
        try:
            tables = pd.read_html(path)
            if tables:
                return tables[0]
        except Exception:
            pass
    try:
        xl = pd.ExcelFile(path)
        best_df = None
        best_score = -1
        for sheet in xl.sheet_names:
            try:
                df = pd.read_excel(path, sheet_name=sheet)
            except Exception:
                continue
            sc = score_dataframe(df)
            if sc > best_score:
                best_df, best_score = df, sc
        if best_df is not None:
            return best_df
    except Exception:
        pass
    try:
        return pd.read_excel(path)
    except Exception:
        pass
    try:
        tables = pd.read_html(path)
        if tables:
            return tables[0]
    except Exception as exc:
        raise RuntimeError(f"Não foi possível ler o arquivo {os.path.basename(path)}. Erro: {exc}") from exc
    raise RuntimeError(f"Não foi possível ler o arquivo {os.path.basename(path)}.")

def find_column(df: pd.DataFrame, expected_names: list[str]) -> str:
    normalized = {normalize_text(col): col for col in df.columns}
    for name in expected_names:
        norm = normalize_text(name)
        if norm in normalized:
            return normalized[norm]
    raise KeyError(f"Coluna não encontrada. Esperado um destes nomes: {expected_names}. Colunas encontradas: {list(df.columns)}")

def prepare_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy().dropna(axis=1, how="all").dropna(axis=0, how="all")
    df.columns = [str(c).strip() for c in df.columns]
    for col in df.columns:
        if normalize_text(col) == "DATA":
            try:
                original = df[col]
                dt = pd.to_datetime(df[col], errors="coerce", dayfirst=True)
                formatted = dt.dt.strftime("%d/%m/%Y")
                df[col] = formatted.where(~formatted.isna(), original.astype(str))
            except Exception:
                pass
    return df

def build_key(name_value, type_value) -> str:
    return normalize_text(name_value) + "||" + normalize_text(type_value)

def build_key_series(name_series: pd.Series, type_series: pd.Series) -> pd.Series:
    return name_series.map(normalize_text) + "||" + type_series.map(normalize_text)

def make_paragraph(text: str, style: ParagraphStyle) -> Paragraph:
    text = "" if pd.isna(text) else str(text)
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
    return Paragraph(text, style)

def build_pdf(df: pd.DataFrame, pdf_path: str, title: str):
    if df.empty:
        raise ValueError("A tabela filtrada ficou vazia. Não há dados para gerar o PDF.")
    page_width, _ = landscape(A4)
    doc = SimpleDocTemplate(pdf_path, pagesize=landscape(A4), leftMargin=10*mm, rightMargin=10*mm, topMargin=10*mm, bottomMargin=10*mm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleCustom", parent=styles["Title"], alignment=TA_CENTER, fontName="Helvetica-Bold", fontSize=12, spaceAfter=6, textColor=colors.black)
    cell_style = ParagraphStyle("Cell", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.0, leading=9.5, alignment=TA_CENTER, spaceAfter=0, spaceBefore=0)
    header_style = ParagraphStyle("Header", parent=cell_style, fontName="Helvetica-Bold", textColor=colors.white)
    headers = [str(col) for col in df.columns]
    data = [[make_paragraph(h, header_style) for h in headers]]
    for _, row in df.iterrows():
        data.append([make_paragraph(row[col], cell_style) for col in df.columns])
    total_width = page_width - doc.leftMargin - doc.rightMargin
    preferred = {"EVENTO":14,"EMPRESA":34,"UNIDADE":30,"NOME":28,"CPF":16,"TIPO":18,"STATUS":18,"DATA":15,"RECIBO ESOCIAL":28,"RECIBO E-SOCIAL":28,"RECIBO SEFAZ":30}
    weights = [preferred.get(normalize_text(col), 18) for col in headers]
    weight_sum = sum(weights)
    col_widths = [total_width*w/weight_sum for w in weights]
    table = Table(data, colWidths=col_widths, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#3A3A3A")),
        ("TEXTCOLOR",(0,0),(-1,0),colors.white),
        ("GRID",(0,0),(-1,-1),0.75,colors.black),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("ALIGN",(0,0),(-1,-1),"CENTER"),
        ("LEFTPADDING",(0,0),(-1,-1),3),
        ("RIGHTPADDING",(0,0),(-1,-1),3),
        ("TOPPADDING",(0,0),(-1,-1),5),
        ("BOTTOMPADDING",(0,0),(-1,-1),5),
    ]))
    doc.build([Paragraph(title, title_style), Spacer(1, 2*mm), table])

def create_output_folder(base_output_dir: str) -> str:
    folder_path = os.path.join(base_output_dir, f"RESULTADO FINAL - {datetime.now().strftime('%Y-%m-%d %H-%M-%S')}")
    os.makedirs(folder_path, exist_ok=True)
    return folder_path

def create_structure(base_folder: str):
    pdf_folder = os.path.join(base_folder, "PDFs")
    log_folder = os.path.join(base_folder, "Logs")
    os.makedirs(pdf_folder, exist_ok=True)
    os.makedirs(log_folder, exist_ok=True)
    return pdf_folder, log_folder

def create_zip_from_folder(folder: str) -> str:
    zip_path = unique_path(folder.rstrip("/\\") + ".zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, _, files in os.walk(folder):
            for filename in files:
                if filename.lower().endswith(".zip"):
                    continue
                full = os.path.join(root, filename)
                zf.write(full, os.path.relpath(full, folder))
    return zip_path

def export_summary_excel(rows: list[dict], excel_path: str):
    wb = Workbook()
    ws = wb.active
    ws.title = "Resumo"
    headers = ["EMPRESA","CNPJ","STATUS","TOTAL BASE EMPRESA","TOTAL ENCONTRADO NO SISTEMA","MOTIVO","PDF GERADO"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)
    for row in rows:
        ws.append([row.get("empresa",""),row.get("cnpj",""),row.get("status",""),row.get("total_base",0),row.get("total_encontrado",0),row.get("motivo",""),row.get("pdf","")])
    for col, width in {"A":45,"B":18,"C":18,"D":20,"E":24,"F":100,"G":55}.items():
        ws.column_dimensions[col].width = width
    wb.save(excel_path)

def get_company_fields_system(system_df: pd.DataFrame):
    normalized_map = {normalize_text(col): col for col in system_df.columns}
    company_col = None
    for wanted in ["EMPRESA", "SETOR", "UNIDADE"]:
        if wanted in normalized_map:
            company_col = normalized_map[wanted]
            break
    if not company_col:
        raise KeyError("Não foi possível identificar a coluna da empresa na planilha do sistema.")
    values = system_df[company_col].dropna().astype(str).str.strip()
    company_text = values.iloc[0] if not values.empty else "EMPRESA"
    return company_col, company_text, extract_cnpj(company_text)

def get_company_fields_base(base_df: pd.DataFrame):
    normalized_map = {normalize_text(col): col for col in base_df.columns}
    for wanted in ["SETOR", "EMPRESA", "UNIDADE"]:
        if wanted in normalized_map:
            return normalized_map[wanted]
    raise KeyError("Não foi possível identificar a coluna da empresa na planilha base.")

def get_best_fuzzy_company_match(work: pd.DataFrame, company_norm: str):
    candidates = work[["__EMPRESA_TXT__","__EMPRESA_NORM__"]].drop_duplicates().to_dict("records")
    best_name = ""
    best_score = 0.0
    for item in candidates:
        candidate = item["__EMPRESA_NORM__"]
        if not candidate:
            continue
        score = SequenceMatcher(None, company_norm, candidate).ratio()
        if company_norm in candidate or candidate in company_norm:
            score += 0.15
        if score > best_score:
            best_name, best_score = item["__EMPRESA_TXT__"], score
    return best_name, best_score

def filter_base_company(base_df: pd.DataFrame, company_text: str, company_cnpj: str, company_col: str):
    work = base_df.copy()
    work["__EMPRESA_TXT__"] = work[company_col].astype(str).str.strip()
    work["__EMPRESA_CNPJ__"] = work["__EMPRESA_TXT__"].map(extract_cnpj)
    work["__EMPRESA_NORM__"] = work["__EMPRESA_TXT__"].map(normalize_company_name)
    company_norm = normalize_company_name(company_text)
    if company_cnpj:
        by_cnpj = work[work["__EMPRESA_CNPJ__"] == company_cnpj].copy()
        if not by_cnpj.empty:
            return by_cnpj, "CNPJ"
    by_exact = work[work["__EMPRESA_NORM__"] == company_norm].copy()
    if not by_exact.empty:
        return by_exact, "NOME EXATO"
    by_contains = work[work["__EMPRESA_NORM__"].str.contains(re.escape(company_norm), na=False, regex=True) | work["__EMPRESA_NORM__"].apply(lambda x: x in company_norm if x else False)].copy()
    if not by_contains.empty:
        return by_contains, "NOME PARCIAL"
    best_name, best_score = get_best_fuzzy_company_match(work, company_norm)
    if best_name and best_score >= 0.75:
        fuzzy = work[work["__EMPRESA_TXT__"] == best_name].copy()
        if not fuzzy.empty:
            return fuzzy, f"NOME APROXIMADO ({best_score:.2f})"
    return work.iloc[0:0].copy(), "NÃO ENCONTRADO"

def reorder_system_by_base(filtered_system: pd.DataFrame, base_company_df: pd.DataFrame, system_nome: str, system_tipo: str, base_nome: str, base_tipo: str):
    filtered_system = filtered_system.copy()
    filtered_system["__CHAVE__"] = build_key_series(filtered_system[system_nome], filtered_system[system_tipo])
    grouped = {}
    for _, row in filtered_system.iterrows():
        grouped.setdefault(row["__CHAVE__"], []).append(row.to_dict())
    ordered_rows = []
    for _, base_row in base_company_df.iterrows():
        key = build_key(base_row[base_nome], base_row[base_tipo])
        for item in grouped.get(key, []):
            ordered_rows.append(item)
    if not ordered_rows:
        return filtered_system[[c for c in filtered_system.columns if c != "__CHAVE__"]].iloc[0:0].copy()
    ordered_df = pd.DataFrame(ordered_rows)
    return ordered_df[[c for c in ordered_df.columns if c != "__CHAVE__"]].reset_index(drop=True)

def run_company_process(system_file: str, base_file: str, pdf_folder: str, log_folder: str, base_sheet: str | None = None):
    system_df = prepare_dataframe(read_spreadsheet(system_file))
    base_df = prepare_dataframe(read_spreadsheet(base_file, selected_sheet=base_sheet))
    system_nome = find_column(system_df, ["NOME"])
    system_tipo = find_column(system_df, ["TIPO"])
    base_nome = find_column(base_df, ["FUNCIONARIO", "FUNCIONÁRIO"])
    base_tipo = find_column(base_df, ["TIPO DE EXAME"])
    base_status = find_column(base_df, ["DEPOSITANTE"])
    _, company_text, company_cnpj = get_company_fields_system(system_df)
    base_company_col = get_company_fields_base(base_df)
    base_company_df, match_method = filter_base_company(base_df, company_text, company_cnpj, base_company_col)
    if base_company_df.empty:
        return {"empresa": company_text, "cnpj": company_cnpj, "status": "NÃO GERADO", "total_base": 0, "total_encontrado": 0, "motivo": "Empresa não encontrada na planilha base.", "pdf": ""}

    base_company_df = base_company_df.copy()
    base_company_df["__STATUS_OK__"] = base_company_df[base_status].map(normalize_text)
    invalid_status = base_company_df[base_company_df["__STATUS_OK__"] != "OK E-SOCIAL"].copy()

    system_df["__CHAVE__"] = build_key_series(system_df[system_nome], system_df[system_tipo])
    base_company_df["__CHAVE__"] = build_key_series(base_company_df[base_nome], base_company_df[base_tipo])

    expected_keys = set(base_company_df["__CHAVE__"].dropna().tolist())
    filtered_system = system_df[system_df["__CHAVE__"].isin(expected_keys)].copy()
    filtered_system = filtered_system[[c for c in system_df.columns if c != "__CHAVE__"]].copy()

    missing_keys = sorted(expected_keys - set(system_df["__CHAVE__"].dropna().tolist()))
    reasons = []
    if not invalid_status.empty:
        for _, row in invalid_status.iterrows():
            reasons.append(f"{row.get(base_nome,'')} | {row.get(base_tipo,'')} | {row.get(base_status,'')}")
    if missing_keys:
        for key in missing_keys:
            name, exam = key.split("||", 1)
            reasons.append(f"NÃO ENCONTRADO NO SISTEMA | {name} | {exam}")

    log_path = unique_path(os.path.join(log_folder, sanitize_filename(company_text) + " - LOG.txt"))
    with open(log_path, "w", encoding="utf-8") as f:
        f.write("RESUMO DO PROCESSAMENTO\n" + "="*80 + "\n")
        f.write(f"Planilha do sistema: {system_file}\nPlanilha base: {base_file}\n")
        f.write(f"Aba base usada: {base_sheet or 'Detecção automática'}\n")
        f.write(f"Empresa do sistema: {company_text}\nCNPJ detectado: {company_cnpj or 'NÃO INFORMADO'}\n")
        f.write(f"Método de correspondência da empresa: {match_method}\n")
        f.write(f"Total base empresa: {len(base_company_df)}\nTotal encontrado no sistema: {len(filtered_system)}\n\n")
        if reasons:
            f.write("MOTIVOS PARA NÃO GERAR PDF\n" + "-"*80 + "\n")
            for item in reasons:
                f.write(f"- {item}\n")
        else:
            f.write("Todos os funcionários da empresa estão com OK E-SOCIAL e foram encontrados no sistema.\n")

    if reasons:
        return {"empresa": company_text, "cnpj": company_cnpj, "status": "NÃO GERADO", "total_base": len(base_company_df), "total_encontrado": len(filtered_system), "motivo": " | ".join(reasons), "pdf": ""}

    ordered_system = reorder_system_by_base(filtered_system, base_company_df, system_nome, system_tipo, base_nome, base_tipo)
    pdf_path = unique_path(os.path.join(pdf_folder, sanitize_filename(company_text) + ".pdf"))
    build_pdf(ordered_system, pdf_path, title=company_text)
    return {"empresa": company_text, "cnpj": company_cnpj, "status": "GERADO", "total_base": len(base_company_df), "total_encontrado": len(ordered_system), "motivo": "OK", "pdf": pdf_path}

# =========================
# FÍSICO E MENTAL
# =========================
def fisico_get_conn():
    conn = sqlite3.connect(FISICO_DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_fisico_db():
    with fisico_get_conn() as conn:
        conn.execute('PRAGMA foreign_keys = ON')
        conn.execute('''
            CREATE TABLE IF NOT EXISTS fisico_empresas (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                nome TEXT NOT NULL UNIQUE,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        conn.execute('''
            CREATE TABLE IF NOT EXISTS fisico_cargos (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                nome TEXT NOT NULL UNIQUE,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        cargo_cols = [row['name'] for row in conn.execute("PRAGMA table_info(fisico_cargos)").fetchall()]
        if 'empresa_id' in cargo_cols:
            conn.execute('ALTER TABLE fisico_cargos RENAME TO fisico_cargos_old')
            conn.execute('''
                CREATE TABLE fisico_cargos (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    nome TEXT NOT NULL UNIQUE,
                    created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            conn.execute('''
                INSERT OR IGNORE INTO fisico_cargos (nome, created_at)
                SELECT DISTINCT nome, COALESCE(created_at, CURRENT_TIMESTAMP)
                FROM fisico_cargos_old
                WHERE nome IS NOT NULL AND TRIM(nome) <> ''
            ''')
            conn.execute('DROP TABLE fisico_cargos_old')
        conn.commit()

def fisico_clean_text(value: str) -> str:
    value = (value or '').strip()
    value = re.sub(r'\s+', ' ', value)
    return value.upper()

def fisico_slugify(value: str) -> str:
    normalized = unicodedata.normalize('NFKD', value or '')
    ascii_text = normalized.encode('ascii', 'ignore').decode('ascii')
    ascii_text = re.sub(r'[^A-Za-z0-9]+', '_', ascii_text).strip('_')
    return ascii_text or 'documento'

def fisico_make_rich(value: str) -> RichText:
    rt = RichText()
    rt.add(fisico_clean_text(value), bold=True)
    return rt

def fisico_build_orgao_texto(empresa: str, edital: str, pss: str) -> str:
    parts = [fisico_clean_text(empresa), fisico_clean_text(edital), fisico_clean_text(pss)]
    parts = [p for p in parts if p]
    return ' - '.join(parts)

def fisico_format_date_extenso(raw_date: str) -> str:
    if raw_date:
        dt = datetime.strptime(raw_date, '%Y-%m-%d').date()
    else:
        dt = datetime.today().date()
    return f'{dt.day:02d} DE {nome_mes(dt.month)} DE {dt.year}'

def fisico_convert_to_pdf(docx_path: str, target_dir: str) -> str:
    soffice = shutil.which('soffice') or shutil.which('libreoffice')
    if not soffice:
        raise RuntimeError('LibreOffice/soffice não encontrado para converter PDF.')
    cmd = [soffice, '--headless', '--convert-to', 'pdf', '--outdir', target_dir, docx_path]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or result.stdout.strip() or 'Falha ao converter para PDF.')
    pdf_path = os.path.join(target_dir, f"{Path(docx_path).stem}.pdf")
    if not os.path.exists(pdf_path):
        raise RuntimeError('PDF não foi gerado.')
    return pdf_path

def fisico_list_empresas(search: str = ''):
    query = 'SELECT id, nome FROM fisico_empresas'
    params = []
    if search:
        query += ' WHERE nome LIKE ?'
        params.append(f'%{search}%')
    query += ' ORDER BY nome'
    with fisico_get_conn() as conn:
        return conn.execute(query, params).fetchall()

def fisico_list_cargos(search: str = ''):
    query = 'SELECT id, nome FROM fisico_cargos'
    params = []
    if search:
        query += ' WHERE nome LIKE ?'
        params.append(f'%{search}%')
    query += ' ORDER BY nome'
    with fisico_get_conn() as conn:
        return conn.execute(query, params).fetchall()

def fisico_render_home(form_data=None):
    form_data = form_data or {}
    return render_template('fisico_mental.html',
                           title='Físico e Mental',
                           today=form_data.get('data_exame') or datetime.today().strftime('%Y-%m-%d'),
                           empresas=fisico_list_empresas(),
                           cargos=fisico_list_cargos(),
                           form_data=form_data)

def fisico_render_cadastros(search=''):
    return render_template('fisico_mental_cadastros.html',
                           title='Cadastros Físico e Mental',
                           search=search,
                           empresas=fisico_list_empresas(search),
                           cargos=fisico_list_cargos(search))

init_fisico_db()

# =========================
# ROTAS
# =========================
@app.route('/fisico-mental', methods=['GET'])
def fisico_mental():
    return fisico_render_home()

@app.route('/fisico-mental/cadastros', methods=['GET'])
def fisico_mental_cadastros():
    search = (request.args.get('q') or '').strip()
    return fisico_render_cadastros(search)

@app.route('/fisico-mental/cadastros/empresa/adicionar', methods=['POST'])
def fisico_adicionar_empresa():
    nome = fisico_clean_text(request.form.get('nome', ''))
    if not nome:
        flash('Digite o nome da empresa.', 'error')
        return redirect(url_for('fisico_mental_cadastros'))
    try:
        with fisico_get_conn() as conn:
            conn.execute('INSERT INTO fisico_empresas (nome) VALUES (?)', (nome,))
            conn.commit()
        flash('Empresa cadastrada com sucesso.', 'success')
    except sqlite3.IntegrityError:
        flash('Essa empresa já está cadastrada.', 'error')
    return redirect(url_for('fisico_mental_cadastros'))

@app.route('/fisico-mental/cadastros/empresa/editar', methods=['POST'])
def fisico_editar_empresa():
    item_id = request.form.get('id', '')
    nome = fisico_clean_text(request.form.get('nome', ''))
    if not item_id.isdigit() or not nome:
        flash('Não foi possível editar a empresa.', 'error')
        return redirect(url_for('fisico_mental_cadastros'))
    try:
        with fisico_get_conn() as conn:
            conn.execute('UPDATE fisico_empresas SET nome = ? WHERE id = ?', (nome, int(item_id)))
            conn.commit()
        flash('Empresa atualizada com sucesso.', 'success')
    except sqlite3.IntegrityError:
        flash('Já existe outra empresa com esse nome.', 'error')
    return redirect(url_for('fisico_mental_cadastros'))

@app.route('/fisico-mental/cadastros/empresa/excluir', methods=['POST'])
def fisico_excluir_empresa():
    item_id = request.form.get('id', '')
    if not item_id.isdigit():
        flash('Não foi possível excluir a empresa.', 'error')
        return redirect(url_for('fisico_mental_cadastros'))
    with fisico_get_conn() as conn:
        conn.execute('PRAGMA foreign_keys = ON')
        conn.execute('DELETE FROM fisico_empresas WHERE id = ?', (int(item_id),))
        conn.commit()
    flash('Empresa excluída.', 'success')
    return redirect(url_for('fisico_mental_cadastros'))

@app.route('/fisico-mental/cadastros/cargo/adicionar', methods=['POST'])
def fisico_adicionar_cargo():
    nome = fisico_clean_text(request.form.get('nome', ''))
    if not nome:
        flash('Digite o nome do cargo.', 'error')
        return redirect(url_for('fisico_mental_cadastros'))
    try:
        with fisico_get_conn() as conn:
            conn.execute('INSERT INTO fisico_cargos (nome) VALUES (?)', (nome,))
            conn.commit()
        flash('Cargo cadastrado com sucesso.', 'success')
    except sqlite3.IntegrityError:
        flash('Esse cargo já está cadastrado.', 'error')
    return redirect(url_for('fisico_mental_cadastros'))

@app.route('/fisico-mental/cadastros/cargo/editar', methods=['POST'])
def fisico_editar_cargo():
    item_id = request.form.get('id', '')
    nome = fisico_clean_text(request.form.get('nome', ''))
    if not item_id.isdigit() or not nome:
        flash('Não foi possível editar o cargo.', 'error')
        return redirect(url_for('fisico_mental_cadastros'))
    try:
        with fisico_get_conn() as conn:
            conn.execute('UPDATE fisico_cargos SET nome = ? WHERE id = ?', (nome, int(item_id)))
            conn.commit()
        flash('Cargo atualizado com sucesso.', 'success')
    except sqlite3.IntegrityError:
        flash('Já existe esse cargo cadastrado.', 'error')
    return redirect(url_for('fisico_mental_cadastros'))

@app.route('/fisico-mental/cadastros/cargo/excluir', methods=['POST'])
def fisico_excluir_cargo():
    item_id = request.form.get('id', '')
    if not item_id.isdigit():
        flash('Não foi possível excluir o cargo.', 'error')
        return redirect(url_for('fisico_mental_cadastros'))
    with fisico_get_conn() as conn:
        conn.execute('DELETE FROM fisico_cargos WHERE id = ?', (int(item_id),))
        conn.commit()
    flash('Cargo excluído.', 'success')
    return redirect(url_for('fisico_mental_cadastros'))

@app.route('/fisico-mental/gerar', methods=['POST'])
def fisico_gerar():
    form_data = request.form.to_dict(flat=True)
    nome = fisico_clean_text(request.form.get('nome', ''))
    rg = fisico_clean_text(request.form.get('rg', ''))
    cpf = fisico_clean_text(request.form.get('cpf', ''))
    empresa = fisico_clean_text(request.form.get('empresa_nome', ''))
    edital = fisico_clean_text(request.form.get('edital', ''))
    pss = fisico_clean_text(request.form.get('pss', ''))
    funcao = fisico_clean_text(request.form.get('funcao_nome', ''))
    data_exame = request.form.get('data_exame', '')
    formato = (request.form.get('formato', 'docx') or 'docx').lower()

    if not nome or not rg or not cpf or not empresa or not funcao:
        flash('Preencha pelo menos: nome, RG, CPF, empresa e cargo.', 'error')
        return fisico_render_home(form_data)

    context = {
        'nome': fisico_make_rich(nome),
        'rg': fisico_make_rich(rg),
        'cpf': fisico_make_rich(cpf),
        'empresa': fisico_make_rich(empresa),
        'edital': fisico_make_rich(edital),
        'pss': fisico_make_rich(pss),
        'orgao_texto': fisico_make_rich(fisico_build_orgao_texto(empresa, edital, pss)),
        'funcao': fisico_make_rich(funcao),
        'data_extenso': fisico_make_rich(fisico_format_date_extenso(data_exame)),
    }

    filename_base = fisico_slugify(f'fisico_mental_{nome}')
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, f'{filename_base}.docx')
        doc = DocxTemplate(FISICO_TEMPLATE_PATH)
        doc.render(context)
        doc.save(docx_path)
        if formato == 'pdf':
            try:
                pdf_path = fisico_convert_to_pdf(docx_path, tmpdir)
                payload = Path(pdf_path).read_bytes()
                return send_file(io.BytesIO(payload), as_attachment=True, download_name=f'{filename_base}.pdf', mimetype='application/pdf')
            except Exception as exc:
                flash(f'Não foi possível gerar PDF agora: {exc}. O arquivo foi enviado em Word.', 'error')
        payload = Path(docx_path).read_bytes()
        return send_file(io.BytesIO(payload), as_attachment=True, download_name=f'{filename_base}.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route("/")
def home():
    return render_template("home.html", title=APP_TITLE)

@app.route("/relatorios", methods=["GET", "POST"])
def relatorios():
    if request.method == "POST":
        files = request.files.getlist("files")
        try:
            mes = int(request.form.get("mes"))
            if mes < 1 or mes > 12:
                raise ValueError
        except Exception:
            flash("Mês inválido.")
            return redirect(url_for("relatorios"))

        arquivos_memoria = [(f.filename, f.read()) for f in files if f and f.filename]
        if not arquivos_memoria:
            flash("Selecione as planilhas.")
            return redirect(url_for("relatorios"))

        files_rel = [UploadedMemoryFile(nome, dados) for nome, dados in arquivos_memoria]
        files_base = [UploadedMemoryFile(nome, dados) for nome, dados in arquivos_memoria]

        wb_rel = criar_relatorio(files_rel, mes)
        wb_base = criar_base(files_base, mes)

        temp_dir = tempfile.mkdtemp()
        caminho_rel = os.path.join(temp_dir, f"Relatorio_{nome_mes(mes)}.xlsx")
        caminho_base = os.path.join(temp_dir, f"Base_do_Mes_{nome_mes(mes)}.xlsx")
        caminho_zip = os.path.join(temp_dir, f"Arquivos_{nome_mes(mes)}.zip")
        wb_rel.save(caminho_rel)
        wb_base.save(caminho_base)

        with zipfile.ZipFile(caminho_zip, "w", zipfile.ZIP_DEFLATED) as z:
            z.write(caminho_rel, os.path.basename(caminho_rel))
            z.write(caminho_base, os.path.basename(caminho_base))

        return send_file(caminho_zip, as_attachment=True, download_name=f"Arquivos_{nome_mes(mes)}.zip")
    return render_template("relatorios.html")

@app.route("/encaminhamentos", methods=["GET", "POST"])
def encaminhamentos():
    if request.method == "POST":
        file = request.files.get("file")
        if not file or not file.filename:
            flash("Envie a planilha Base do Mês.")
            return redirect(url_for("encaminhamentos"))
        zip_path = gerar_encaminhamentos(file)
        return send_file(zip_path, as_attachment=True, download_name="encaminhamentos.zip")
    return render_template("encaminhamentos.html")

@app.route("/renumerador", methods=["GET", "POST"])
def renumerador():
    if request.method == "POST":
        arquivo = request.files.get("arquivo")
        nova_data = request.form.get("nova_data", "").strip()

        if not arquivo or not arquivo.filename:
            flash("Selecione um arquivo.")
            return redirect(url_for("renumerador"))

        if not nova_data:
            flash("Informe a data.")
            return redirect(url_for("renumerador"))

        if not allowed_renum_file(arquivo.filename):
            flash("Formato inválido. Envie um arquivo .docx ou .zip.")
            return redirect(url_for("renumerador"))

        with tempfile.TemporaryDirectory() as temp_dir_str:
            temp_dir = Path(temp_dir_str)
            entrada_dir = temp_dir / "entrada"
            saida_dir = temp_dir / "saida"
            entrada_dir.mkdir(parents=True, exist_ok=True)
            saida_dir.mkdir(parents=True, exist_ok=True)

            nome_seguro = secure_filename(arquivo.filename)
            caminho_upload = temp_dir / nome_seguro
            arquivo.save(caminho_upload)

            if caminho_upload.suffix.lower() == ".zip":
                try:
                    with zipfile.ZipFile(caminho_upload, "r") as zip_ref:
                        zip_ref.extractall(entrada_dir)
                except zipfile.BadZipFile:
                    flash("O arquivo ZIP enviado está corrompido ou inválido.")
                    return redirect(url_for("renumerador"))
            else:
                destino = entrada_dir / nome_seguro
                destino.write_bytes(caminho_upload.read_bytes())

            arquivos_docx = [p for p in entrada_dir.rglob("*.docx") if not p.name.startswith("~$")]
            if not arquivos_docx:
                flash("Nenhum arquivo .docx foi encontrado para processar.")
                return redirect(url_for("renumerador"))

            total_arquivos = 0
            total_recibos = 0
            relatorio = []

            for caminho in arquivos_docx:
                relativo = caminho.relative_to(entrada_dir)
                destino = saida_dir / relativo
                try:
                    alterados, ultimo, datas_alteradas = renumerar_documento(str(caminho), str(destino), nova_data)
                    total_arquivos += 1
                    total_recibos += alterados
                    relatorio.append(f"{relativo.as_posix()} | último encontrado: {ultimo} | recibos renumerados: {alterados} | datas alteradas: {datas_alteradas}")
                except Exception as e:
                    relatorio.append(f"{relativo.as_posix()} | erro: {str(e)}")

            relatorio_path = saida_dir / "relatorio_processamento.txt"
            relatorio_path.write_text(
                "Renumerador de Recibos - Relatório de Processamento\n\n"
                f"Arquivos processados: {total_arquivos}\n"
                f"Recibos renumerados: {total_recibos}\n\n" + "\n".join(relatorio),
                encoding="utf-8"
            )

            if len(arquivos_docx) == 1:
                unico_saida = next((p for p in saida_dir.rglob("*.docx")), None)
                if unico_saida:
                    buffer = io.BytesIO(unico_saida.read_bytes())
                    buffer.seek(0)
                    return send_file(buffer, as_attachment=True, download_name=unico_saida.name, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            zip_path = temp_dir / "recibos_renumerados.zip"
            with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zip_out:
                for arquivo_saida in saida_dir.rglob("*"):
                    if arquivo_saida.is_file():
                        zip_out.write(arquivo_saida, arquivo_saida.relative_to(saida_dir))

            zip_buffer = io.BytesIO(zip_path.read_bytes())
            zip_buffer.seek(0)
            return send_file(zip_buffer, as_attachment=True, download_name="recibos_renumerados.zip", mimetype="application/zip")

    return render_template("renumerador.html")

@app.route("/esocial", methods=["GET"])
def esocial():
    return render_template("esocial.html", title="E-SOCIAL EVELLYN")

@app.route("/esocial/abas-base", methods=["POST"])
def esocial_abas_base():
    base_file = request.files.get("base_file")
    if not base_file or not base_file.filename:
        return jsonify({"ok": False, "error": "Nenhuma planilha base enviada."}), 400
    temp_root = Path(tempfile.mkdtemp(prefix="esocial_abas_"))
    try:
        base_path = temp_root / secure_filename(base_file.filename)
        base_file.save(base_path)
        return jsonify({"ok": True, "sheets": list_sheets(str(base_path))})
    except Exception as exc:
        return jsonify({"ok": False, "error": str(exc)}), 500

@app.route("/esocial/processar", methods=["POST"])
def esocial_processar():
    base_file = request.files.get("base_file")
    rel_files = request.files.getlist("rel_files")
    base_sheet = request.form.get("base_sheet", "").strip() or None

    if not base_file or not base_file.filename:
        flash("Selecione a planilha base.")
        return redirect(url_for("esocial"))

    valid_rel_files = [f for f in rel_files if f and f.filename and is_allowed_file(f.filename)]
    if not valid_rel_files:
        flash("Selecione um ou mais arquivos RELFUNCGERAL válidos.")
        return redirect(url_for("esocial"))

    temp_root = Path(tempfile.mkdtemp(prefix="esocial_web_"))
    upload_dir = temp_root / "uploads"
    output_root = temp_root / "saida"
    upload_dir.mkdir(parents=True, exist_ok=True)
    output_root.mkdir(parents=True, exist_ok=True)

    try:
        base_path = upload_dir / secure_filename(base_file.filename)
        base_file.save(base_path)

        rel_paths = []
        for index, rel in enumerate(valid_rel_files, start=1):
            filename = secure_filename(Path(rel.filename).name)
            path = upload_dir / f"{index:03d}_{filename}"
            rel.save(path)
            rel_paths.append(path)

        general_output_folder = Path(create_output_folder(str(output_root)))
        pdf_folder, log_folder = create_structure(str(general_output_folder))
        summary_rows = []

        for rel_path in rel_paths:
            try:
                summary_rows.append(run_company_process(str(rel_path), str(base_path), pdf_folder, log_folder, base_sheet=base_sheet))
            except Exception as exc:
                summary_rows.append({"empresa": rel_path.name, "cnpj": "", "status": "NÃO GERADO", "total_base": 0, "total_encontrado": 0, "motivo": str(exc), "pdf": ""})

        resumo_excel = str(Path(general_output_folder) / "RESUMO_FINAL.xlsx")
        export_summary_excel(summary_rows, resumo_excel)

        with open(Path(log_folder) / "RESUMO_GERAL.txt", "w", encoding="utf-8") as f:
            f.write("RESUMO GERAL DO PROCESSAMENTO\n" + "="*80 + "\n\n")
            f.write(f"Planilha base: {base_path.name}\nAba base usada: {base_sheet or 'Detecção automática'}\n")
            f.write(f"Total de empresas processadas: {len(summary_rows)}\n")
            f.write(f"PDFs gerados: {sum(1 for r in summary_rows if r['status']=='GERADO')}\n")
            f.write(f"PDFs não gerados: {sum(1 for r in summary_rows if r['status']!='GERADO')}\n")
            f.write(f"Resumo Excel: {resumo_excel}\n")

        zip_path = Path(create_zip_from_folder(str(general_output_folder)))
        return send_file(zip_path, as_attachment=True, download_name=zip_path.name, mimetype="application/zip")
    except Exception as exc:
        flash(f"Erro ao processar: {exc}")
        return redirect(url_for("esocial"))

if __name__ == "__main__":
    init_fisico_db()
    app.run(debug=False, host="0.0.0.0", port=5000)
